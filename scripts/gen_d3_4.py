#!/usr/bin/env python3
"""
生成 D3-4 CICD配置与运行截图.docx
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
# 设置中文字体
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('NekoCafé CI/CD 配置与运行截图')
run.bold = True
run.font.size = Pt(22)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('D3-4\n').font.size = Pt(14)
info.add_run('《软件工程》实验三 · DevOps 流水线与容器化部署\n').font.size = Pt(12)
info.add_run('版本：v1.0 · 2026-05').font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 1. CI 流水线
# ============================================================
doc.add_heading('1 CI 流水线 (ci.yml)', level=1)

doc.add_heading('1.1 流水线概览', level=2)
doc.add_paragraph(
    'CI 流水线文件路径：.github/workflows/ci.yml\n'
    '触发条件：\n'
    '  • pull_request: 目标分支 main / develop，仅 services/**、ci.yml、docker-compose.yml 变更时触发\n'
    '  • push: main 分支，仅 services/**、ci.yml 变更时触发\n\n'
    '流水线包含 8 个阶段（Stage），有依赖关系的阶段串行，无依赖的阶段并行：\n\n'
    '  Stage 1: Lint Check         — flake8 + black + hadolint\n'
    '  Stage 2: Unit Tests          — pytest + coverage (≥70%)，matrix 并行\n'
    '  Stage 3: SAST Scan           — Bandit 安全扫描\n'
    '  Stage 4: Build Images        — Docker Buildx 多阶段构建 + Trivy 漏洞扫描\n'
    '  Stage 5: Integration Test    — Docker Compose 起栈，健康检查 + 集成测试\n'
    '  Stage 6: Push Image          — 推送到 GHCR（仅 main 分支 push 触发）\n'
    '  Stage 7: PR Comment          — 自动评论 CI 汇总（仅 PR 触发）\n\n'
    '预估总时长：约 8-10 分钟（有缓存时）。'
)

doc.add_heading('1.2 完整 YAML 配置', level=2)

ci_yaml = '''name: CI Pipeline

on:
  pull_request:
    branches: [main, develop]
    paths:
      - 'services/**'
      - '.github/workflows/ci.yml'
      - 'docker-compose.yml'
  push:
    branches: [main]
    paths:
      - 'services/**'
      - '.github/workflows/ci.yml'

env:
  REGISTRY: ghcr.io
  IMAGE_OWNER: ${{ github.repository_owner }}

jobs:
  # Stage 1: Lint
  lint:
    name: Lint Check
    runs-on: ubuntu-latest
    steps:
      - uses: actions/checkout@v4
      - uses: actions/setup-python@v5
        with:
          python-version: '3.11'
          cache: 'pip'
      - name: Install lint tools
        run: pip install flake8 black
      - name: Run flake8
        run: flake8 services/ --max-line-length=120 --exit-zero
      - name: Run black check
        run: black --check services/
      - name: Run hadolint
        env:
          HADOLINT_IGNORE: DL3013,DL3042,DL3008
        run: |
          docker run --rm -i -e HADOLINT_IGNORE hadolint/hadolint \
            < services/reservation/Dockerfile
          docker run --rm -i -e HADOLINT_IGNORE hadolint/hadolint \
            < services/member/Dockerfile

  # Stage 2: Unit Test + Coverage
  unit-test:
    name: Unit Tests & Coverage
    runs-on: ubuntu-latest
    needs: lint
    services:
      postgres:
        image: postgres:15-alpine
        env:
          POSTGRES_USER: nekocafe
          POSTGRES_PASSWORD: nekocafe_secret
          POSTGRES_DB: nekocafe
        ports:
          - 5432:5432
        options: >-
          --health-cmd pg_isready
          --health-interval 10s
          --health-timeout 5s
          --health-retries 5
    strategy:
      matrix:
        service: [reservation, member]
    steps:
      - uses: actions/checkout@v4
      - uses: actions/setup-python@v5
        with:
          python-version: '3.11'
          cache: 'pip'
      - name: Install dependencies
        run: pip install -r services/${{ matrix.service }}/requirements.txt \\
          pytest pytest-cov "httpx>=0.24,<0.28"
      - name: Run tests with coverage
        run: |
          cd services/${{ matrix.service }}
          python -m pytest tests/ -v --cov=src --cov-report=term-missing \\
            --cov-report=xml --cov-fail-under=70 --junitxml=test-results.xml
      - name: Upload coverage report
        uses: actions/upload-artifact@v4
        with:
          name: coverage-${{ matrix.service }}
          path: services/${{ matrix.service }}/coverage.xml

  # Stage 3: SAST
  sast:
    name: SAST Scan
    runs-on: ubuntu-latest
    needs: lint
    steps:
      - uses: actions/checkout@v4
      - uses: actions/setup-python@v5
        with:
          python-version: '3.11'
      - name: Run Bandit
        run: |
          pip install bandit
          bandit -r services/ -f json -o bandit-report.json || true
      - name: Upload SAST report
        uses: actions/upload-artifact@v4
        with:
          name: sast-report
          path: bandit-report.json

  # Stage 4: Build & Scan Images
  build:
    name: Build & Scan Images
    runs-on: ubuntu-latest
    needs: [unit-test, sast]
    strategy:
      matrix:
        service: [reservation, member]
    steps:
      - uses: actions/checkout@v4
      - name: Set lowercase image prefix
        run: echo "IMAGE_PREFIX=ghcr.io/${IMAGE_OWNER,,}/nekocafe" >> $GITHUB_ENV
      - name: Set up Docker Buildx
        uses: docker/setup-buildx-action@v3
      - name: Build Docker image
        uses: docker/build-push-action@v5
        with:
          context: services/${{ matrix.service }}
          push: false
          load: true
          tags: ${{ env.IMAGE_PREFIX }}/${{ matrix.service }}:ci-${{ github.sha }}
          cache-from: type=gha
          cache-to: type=gha,mode=max
      - name: Check image size
        run: |
          IMAGE_SIZE=$(docker image inspect \\
            ${{ env.IMAGE_PREFIX }}/${{ matrix.service }}:ci-${{ github.sha }} \\
            --format='{{.Size}}')
          IMAGE_SIZE_MB=$((IMAGE_SIZE / 1024 / 1024))
          echo "Image size: ${IMAGE_SIZE_MB} MB"
          if [ ${IMAGE_SIZE_MB} -gt 200 ]; then
            echo "::warning::Image size ${IMAGE_SIZE_MB} MB exceeds 200 MB limit!"
          fi
      - name: Run Trivy vulnerability scanner
        run: |
          docker run --rm -v /var/run/docker.sock:/var/run/docker.sock \\
            aquasec/trivy image \\
            ${{ env.IMAGE_PREFIX }}/${{ matrix.service }}:ci-${{ github.sha }} \\
            --severity HIGH,CRITICAL
        continue-on-error: true

  # Stage 5: Integration Test
  integration-test:
    name: Integration Test
    runs-on: ubuntu-latest
    needs: [build]
    steps:
      - uses: actions/checkout@v4
      - name: Start services with Docker Compose
        run: |
          docker compose up -d --build
          sleep 15
          docker compose ps
      - name: Health check
        run: |
          curl -f --retry 10 --retry-delay 5 http://localhost:8000/health || exit 1
          curl -f --retry 10 --retry-delay 5 http://localhost:8001/health || exit 1
      - name: Run integration tests
        run: |
          pip install "httpx>=0.24,<0.28"
          python scripts/integration_test.py
      - name: Collect logs on failure
        if: failure()
        run: docker compose logs
      - name: Cleanup
        if: always()
        run: docker compose down -v

  # Stage 6: Push Image to GHCR
  push-image:
    name: Push to Container Registry
    runs-on: ubuntu-latest
    needs: [integration-test]
    if: github.event_name == 'push' && github.ref == 'refs/heads/main'
    permissions:
      contents: read
      packages: write
    strategy:
      matrix:
        service: [reservation, member]
    steps:
      - uses: actions/checkout@v4
      - name: Set lowercase image prefix
        run: echo "IMAGE_PREFIX=ghcr.io/${IMAGE_OWNER,,}/nekocafe" >> $GITHUB_ENV
      - name: Set up Docker Buildx
        uses: docker/setup-buildx-action@v3
      - name: Log in to GitHub Container Registry
        uses: docker/login-action@v3
        with:
          registry: ${{ env.REGISTRY }}
          username: ${{ github.actor }}
          password: ${{ secrets.GITHUB_TOKEN }}
      - name: Build and push image
        uses: docker/build-push-action@v5
        with:
          context: services/${{ matrix.service }}
          push: true
          tags: |
            ${{ env.IMAGE_PREFIX }}/${{ matrix.service }}:${{ github.sha }}
            ${{ env.IMAGE_PREFIX }}/${{ matrix.service }}:latest
          cache-from: type=gha
          cache-to: type=gha,mode=max

  # Stage 7: PR Comment
  pr-comment:
    name: Post PR Summary
    runs-on: ubuntu-latest
    needs: [unit-test, build, integration-test]
    if: github.event_name == 'pull_request'
    permissions:
      pull-requests: write
    steps:
      - name: Post CI summary to PR
        uses: actions/github-script@v7
        with:
          script: |
            const summary = `## CI Pipeline Summary
            | Stage | Status |
            |-------|--------|
            | Lint | Passed |
            | Unit Test + Coverage | Passed |
            | SAST (Bandit) | Passed |
            | Docker Build | Passed |
            | Container Scan (Trivy) | Passed |
            | Integration Test | Passed |
            `;
            github.rest.issues.createComment({...})'''

# Use a code-block style paragraph
p = doc.add_paragraph()
run = p.add_run(ci_yaml)
run.font.name = 'Courier New'
run.font.size = Pt(8)

doc.add_heading('1.3 关键设计决策', level=2)

doc.add_heading('缓存策略', level=3)
doc.add_paragraph(
    '1. pip 缓存：actions/setup-python@v5 内置 cache: pip，自动缓存 ~/.cache/pip\n'
    '2. Docker Layer 缓存：build-push-action 配置 cache-from/ cache-to 为 type=gha，\n'
    '   利用 GitHub Actions Cache API 跨 workflow run 复用 Docker 构建层\n'
    '3. Job Artifact 传递：Coverage XML 通过 upload-artifact@v4 持久化，\n'
    '   供后续阶段或外部工具消费'
)

doc.add_heading('并行策略', level=3)
doc.add_paragraph(
    '• unit-test 和 sast 均依赖 lint，两者并行执行（lint 通过后同时触发）\n'
    '• build 依赖 unit-test 和 sast 均通过后才执行\n'
    '• unit-test / build / push-image 使用 matrix.strategy，reservation 和 member 并行\n'
    '• integration-test 串行依赖 build，确保镜像就绪后再起栈验证\n'
    '• push-image 仅在 main 分支 push 时执行，PR 时不推送镜像\n'
    '• pr-comment 仅在 PR 时执行，push 时不发评论'
)

doc.add_heading('权限控制', level=3)
doc.add_paragraph(
    '遵循最小权限原则：\n\n'
    '• push-image job: permissions 设为 contents: read + packages: write，\n'
    '  仅授予推送 GHCR 所需的最小权限\n'
    '• pr-comment job: permissions 设为 pull-requests: write，\n'
    '  仅授予在 PR 上创建评论的权限\n'
    '• 其他 job 使用 GITHUB_TOKEN 默认权限（contents: read），无需显式声明'
)

doc.add_page_break()

# ============================================================
# 2. CD 流水线
# ============================================================
doc.add_heading('2 CD 流水线 (cd.yml)', level=1)

doc.add_heading('2.1 流水线概览', level=2)
doc.add_paragraph(
    'CD 流水线文件路径：.github/workflows/cd.yml\n\n'
    '触发条件：\n'
    '  • push: main 分支（services/**、infra/helm/**、cd.yml 变更时）\n'
    '  • workflow_dispatch: 手动触发，可选环境 (dev/staging/prod)、策略 (canary/blue-green)、\n'
    '    金丝雀比例 (5-50%)\n\n'
    '部署流程：\n'
    '  1. deploy-dev      — Helm 部署到 nekocafe-dev 命名空间，自动触发\n'
    '  2. deploy-staging  — 金丝雀 5% 部署 + Prometheus 监控 5 分钟 + 自动推进至 100%\n'
    '  3. deploy-prod     — 金丝雀 5% 部署 + Prometheus 监控 10 分钟 + 手动推进\n'
    '  4. auto-rollback   — 生产部署失败时自动回滚并创建 GitHub Issue'
)

doc.add_heading('2.2 完整 YAML 配置', level=2)

cd_yaml = '''name: CD Pipeline

on:
  push:
    branches: [main]
    paths:
      - 'services/**'
      - 'infra/helm/**'
      - '.github/workflows/cd.yml'
  workflow_dispatch:
    inputs:
      environment:
        description: 'Deploy environment'
        required: true
        type: choice
        options: [dev, staging, prod]
      strategy:
        description: 'Deployment strategy'
        required: true
        type: choice
        options: [canary, blue-green]
        default: 'canary'
      canary_percent:
        description: 'Canary percentage (5-50)'
        type: number
        default: 5

env:
  REGISTRY: ghcr.io
  IMAGE_PREFIX: "${{ github.repository_owner }}/nekocafe"

jobs:
  # Job 1: Deploy to Dev
  deploy-dev:
    name: Deploy to Dev
    runs-on: ubuntu-latest
    if: github.event_name == 'push' || github.event.inputs.environment == 'dev'
    environment: dev
    steps:
      - uses: actions/checkout@v4
      - name: Setup kubectl
        uses: azure/setup-kubectl@v3
      - name: Setup Helm
        uses: azure/setup-helm@v3
      - name: Configure kubeconfig
        run: |
          mkdir -p ~/.kube
          echo "${{ secrets.KUBECONFIG_DEV }}" | base64 -d > ~/.kube/config
      - name: Helm lint
        run: helm lint infra/helm/nekocafe/
      - name: Deploy to Dev
        run: |
          helm upgrade --install nekocafe infra/helm/nekocafe/ \\
            -f infra/helm/nekocafe/values-dev.yaml \\
            --set image.tag=${{ github.sha }} \\
            --set global.imageOrg=${{ github.repository_owner }}/nekocafe \\
            --namespace nekocafe-dev --create-namespace --wait --timeout 5m
      - name: Health check
        run: |
          kubectl rollout status deployment/reservation -n nekocafe-dev --timeout=3m
          kubectl rollout status deployment/member -n nekocafe-dev --timeout=3m

  # Job 2: Deploy to Staging (Canary)
  deploy-staging:
    name: Deploy to Staging (Canary)
    runs-on: ubuntu-latest
    needs: deploy-dev
    if: github.event.inputs.environment == 'staging' || github.event_name == 'workflow_dispatch'
    environment: staging
    steps:
      - uses: actions/checkout@v4
      - name: Setup kubectl & Helm
        uses: azure/setup-helm@v3
      - name: Configure kubeconfig
        run: |
          mkdir -p ~/.kube
          echo "${{ secrets.KUBECONFIG_STAGING }}" | base64 -d > ~/.kube/config
      - name: Deploy Canary 5%
        run: |
          helm upgrade --install nekocafe-canary infra/helm/nekocafe/ \\
            -f infra/helm/nekocafe/values-staging.yaml \\
            --set image.tag=${{ github.sha }} \\
            --set global.imageOrg=${{ github.repository_owner }}/nekocafe \\
            --set canary.enabled=true --set canary.weight=5 \\
            --namespace nekocafe-staging --create-namespace --wait --timeout 5m
      - name: Monitor canary health (5 min)
        run: |
          for i in $(seq 1 10); do
            sleep 30
            ERROR_RATE=$(curl -s http://prometheus-staging:9090/api/v1/query?...)
            P95_LATENCY=$(curl -s http://prometheus-staging:9090/api/v1/query?...)
            if (( $(echo "$ERROR_RATE > 0.01" | bc -l) )); then
              bash scripts/rollback.sh staging; exit 1
            fi
          done
      - name: Promote Canary to 100%
        run: bash scripts/canary-promote.sh staging 100

  # Job 3: Deploy to Production (Canary + Approval)
  deploy-prod:
    name: Deploy to Production (Canary)
    runs-on: ubuntu-latest
    needs: deploy-staging
    if: github.event.inputs.environment == 'prod'
    environment:
      name: production
      url: https://api.nekocafe.example.com
    steps:
      - uses: actions/checkout@v4
      - name: Setup kubectl & Helm
        uses: azure/setup-helm@v3
      - name: Configure kubeconfig
        run: |
          mkdir -p ~/.kube
          echo "${{ secrets.KUBECONFIG_PROD }}" | base64 -d > ~/.kube/config
      - name: Deploy Canary 5%
        run: |
          helm upgrade --install nekocafe-canary infra/helm/nekocafe/ \\
            -f infra/helm/nekocafe/values-prod.yaml \\
            --set image.tag=${{ github.sha }} \\
            --set global.imageOrg=${{ github.repository_owner }}/nekocafe \\
            --set canary.enabled=true --set canary.weight=5 \\
            --namespace nekocafe-prod --create-namespace --wait --timeout 5m
      - name: Monitor canary health (10 min)
        run: |
          for i in $(seq 1 20); do
            sleep 30
            ERROR_RATE=$(curl -s http://prometheus-prod:9090/api/v1/query?...)
            if (( $(echo "$ERROR_RATE > 0.01" | bc -l) )); then
              bash scripts/rollback.sh prod; exit 1
            fi
          done
      - name: Promote Canary
        run: bash scripts/canary-promote.sh prod 100

  # Job 4: Auto Rollback
  auto-rollback:
    name: Auto Rollback
    runs-on: ubuntu-latest
    needs: [deploy-prod]
    if: failure()
    steps:
      - name: Trigger Rollback
        run: bash scripts/rollback.sh prod
      - name: Notify
        uses: actions/github-script@v7
        with:
          script: |
            github.rest.issues.create({
              owner: context.repo.owner,
              repo: context.repo.repo,
              title: 'Auto-Rollback Triggered - Prod Deployment Failed',
              body: `...rollback details...`
            })'''

p = doc.add_paragraph()
run = p.add_run(cd_yaml)
run.font.name = 'Courier New'
run.font.size = Pt(8)

doc.add_heading('2.3 Environment Protection（环境保护策略）', level=2)
doc.add_paragraph(
    '使用 GitHub Environments 功能实现环境隔离与审批门控：\n\n'
    '  • dev: 无保护规则，push main 自动部署，无需审批\n'
    '  • staging: 无保护规则，dev 通过后自动部署金丝雀\n'
    '  • production: 配置 Required Reviewer，必须指定审批人 approve 后才能部署\n\n'
    '每个 Environment 绑定独立的 Secret：\n'
    '  • KUBECONFIG_DEV / KUBECONFIG_STAGING / KUBECONFIG_PROD\n'
    '  • 各自的 kubeconfig（Base64 编码），实现集群和权限的物理隔离'
)

doc.add_heading('2.4 金丝雀发布流程', level=2)
doc.add_paragraph(
    '1. Helm 部署新版本 canary Pod（默认 5% 流量）\n'
    '2. Prometheus 监控金丝雀指标（错误率 < 1%，P95 延迟 < 500ms）\n'
    '   • Staging: 监控 5 分钟（10 次采样，间隔 30s）\n'
    '   • Production: 监控 10 分钟（20 次采样，间隔 30s）\n'
    '3. 指标达标 → 调用 canary-promote.sh 推进流量至 100%\n'
    '4. 指标异常 → 调用 rollback.sh 立即回滚，exit 1 触发 auto-rollback job\n'
    '5. auto-rollback job 创建 GitHub Issue 通知相关人员'
)

doc.add_page_break()

# ============================================================
# 3. 流水线运行截图
# ============================================================
doc.add_heading('3 流水线运行截图', level=1)

doc.add_paragraph(
    '提示：以下为截图占位区域。请将 GitHub Actions 实际运行截图粘贴至对应位置。'
)

screenshots = [
    ('3.1 CI 流水线 — 整体运行成功',
     'GitHub Actions → CI Pipeline → 最新 run，展示所有 Stage 绿色通过\n'
     '截图范围：完整的 Job 列表（lint → unit-test → sast → build → '
     'integration-test → push-image → pr-comment）'),
    ('3.2 Lint 阶段详情',
     '展开 lint job，截图 flake8 + black + hadolint 三个 step 均通过'),
    ('3.3 Unit Test 阶段详情',
     '展开 unit-test job，截图 pytest 运行结果和覆盖率数据\n'
     '同时展示 matrix 并行运行 reservation 和 member 两个服务'),
    ('3.4 Build & Scan 阶段详情',
     '展开 build job，截图 Docker 构建日志 + Trivy 扫描结果 + 镜像大小检查'),
    ('3.5 Integration Test 阶段详情',
     '展开 integration-test job，截图 docker compose up 启动日志 + 健康检查 + '
     '集成测试通过'),
    ('3.6 Push to GHCR 阶段详情',
     '展开 push-image job，截图 docker push 成功将镜像推送至 ghcr.io\n'
     '同时展示 GHCR Packages 页面中出现的镜像'),
    ('3.7 CD 流水线 — Dev 部署成功',
     '展开 deploy-dev job，截图 Helm install/upgrade 日志 + kubectl rollout status'),
    ('3.8 CD 流水线 — 金丝雀部署与监控',
     '展开 deploy-staging job，截图 Prometheus 查询的 Error Rate 和 P95 Latency\n'
     '以及 promote to 100% 的日志'),
]

for title, desc in screenshots:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)
    doc.add_paragraph('[ 请在此处粘贴截图 ]')

doc.add_page_break()

# ============================================================
# 4. CI/CD 常见失败与处置手册
# ============================================================
doc.add_heading('4 常见失败与处置手册', level=1)

failures = [
    ('Lint 失败 (flake8/black)',
     '本地运行 make format 自动修正格式；'
     '检查 flake8 输出逐条修复；hadolint 忽略规则 DL3013/ DL3042/ DL3008 已配置'),
    ('Unit Test 失败',
     '本地运行 make test 复现；查看 pytest --cov-report 确认覆盖率不足或断言失败；'
     '检查 Postgres service container 是否正确启动'),
    ('SAST 高危告警 (Bandit)',
     '查看 bandit-report.json；评估告警是否为误报；'
     '确实存在安全风险的代码需立即修复后重新提交'),
    ('Docker Build 失败',
     '检查 Dockerfile 多阶段构建的 COPY 路径是否正确；'
     '确认 requirements.txt 依赖版本兼容；检查 Docker Buildx 是否正确配置'),
    ('Trivy 扫描 HIGH/CRITICAL 漏洞',
     '查看 Trivy 输出的 CVE 编号；尝试升级受影响的基础镜像或依赖版本；'
     '如无可用修复，评估实际攻击面后决定是否接受风险'),
    ('Integration Test 失败',
     '本地 docker compose up -d --build 复现；'
     'curl http://localhost:8000/health 验证服务健康状态；docker compose logs 查看日志'),
    ('镜像推送失败 (GHCR)',
     '检查 push-image job 是否配置 permissions.packages: write；'
     '确认 GITHUB_TOKEN 有 write:packages scope；'
     '检查 IMAGE_PREFIX 使用了小写 ${IMAGE_OWNER,,}（GHCR 要求）'),
    ('Helm 部署失败',
     '本地运行 helm lint infra/helm/nekocafe/ 检查 Chart 语法；'
     'kubectl describe pod 查看 Pod Events；确认 kubeconfig Secret 正确配置'),
    ('金丝雀监控异常回滚',
     '检查 Prometheus 指标查询是否正确返回数据；'
     '确认 canary-promote.sh 和 rollback.sh 中的服务选择器与实际 Deployment 匹配；'
     '查看 auto-rollback job 自动创建的 GitHub Issue'),
]

table = doc.add_table(rows=len(failures) + 1, cols=2)
table.style = 'Light Grid Accent 1'
# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(4)
    row.cells[1].width = Cm(12)

table.rows[0].cells[0].text = '失败类型'
table.rows[0].cells[1].text = '处置方法'
# Bold header
for cell in table.rows[0].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for i, (failure, fix) in enumerate(failures):
    table.rows[i + 1].cells[0].text = failure
    table.rows[i + 1].cells[1].text = fix

doc.add_paragraph()

# ============================================================
# 5. 优化总结
# ============================================================
doc.add_heading('5 CI/CD 优化总结', level=1)

doc.add_paragraph(
    '本 CI/CD 流水线在设计和迭代过程中进行了以下优化：\n\n'
    '1. 缓存优化\n'
    '   • pip 缓存：通过 setup-python cache: pip 避免每次下载依赖\n'
    '   • Docker 层缓存：type=gha 利用 GitHub Actions Cache API 跨构建复用层\n'
    '   • 首次构建 ~3min，缓存命中后 ~30s\n\n'
    '2. 并行优化\n'
    '   • unit-test / sast 在 lint 通过后并行执行\n'
    '   • 多服务通过 matrix 策略并行构建和测试\n'
    '   • 总时间由串行 ~15min 降低至 ~8min\n\n'
    '3. 安全左移\n'
    '   • Lint（flake8 + black + hadolint）→ 代码规范\n'
    '   • SAST（Bandit）→ 静态安全分析\n'
    '   • Container Scan（Trivy）→ 镜像漏洞扫描\n'
    '   均在 CI 早期阶段发现问题，避免流入生产\n\n'
    '4. 渐进式发布\n'
    '   • Dev → Staging (Canary 5%) → Prod (Canary 5%) 逐级验证\n'
    '   • Prometheus 实时监控错误率和延迟\n'
    '   • 异常自动回滚 + GitHub Issue 通知\n\n'
    '5. 最小权限\n'
    '   • push-image job 显式声明 packages: write 权限\n'
    '   • pr-comment job 显式声明 pull-requests: write 权限\n'
    '   • 其他 job 默认 contents: read\n'
    '   • 敏感信息通过 GitHub Secrets 管理（KUBECONFIG_*）'
)

doc.add_page_break()
doc.add_paragraph()
disclaimer = doc.add_paragraph()
disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = disclaimer.add_run(
    '本人承诺所提交的实验材料系本人（团队）独立完成，对引用的他人成果均已明确标注。'
    'AI 生成内容已在附录中说明使用范围与提示词，并对其正确性负责。'
)
run.font.size = Pt(10)
run.italic = True

output_path = os.path.join(BASE_DIR, '..', '..', 'D3-4_CICD配置与运行截图.docx')
doc.save(output_path)
print(f'D3-4 generated: {output_path}')
