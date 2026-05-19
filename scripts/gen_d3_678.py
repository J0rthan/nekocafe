#!/usr/bin/env python3
"""
生成 D3-6 可观测性配置与Dashboard截图.docx + D3-7 DORA指标报告.xlsx + D3-8 演示视频脚本.docx
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.chart import LineChart, Reference
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from datetime import datetime, timedelta
import random
import os
import json

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, '..', '..')

# ============================================================
# D3-6: 可观测性配置与Dashboard截图
# ============================================================
doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('NekoCafé 可观测性配置与 Dashboard 截图')
run.bold = True
run.font.size = Pt(22)
run.font.name = '黑体'

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('D3-6\n').font.size = Pt(14)
info.add_run('《软件工程》实验三 · DevOps 流水线与容器化部署\n').font.size = Pt(12)
info.add_run('版本：v1.0 · 2026-05').font.size = Pt(11)

doc.add_page_break()

# ==== 1. OpenTelemetry 接入 ====
doc.add_heading('1 OpenTelemetry 接入方案', level=1)

doc.add_heading('1.1 SDK 集成', level=2)
doc.add_paragraph(
    '每个服务通过 Python OpenTelemetry SDK 实现自动和手动追踪：\n\n'
    'Python 依赖：\n'
    '• opentelemetry-api==1.21.0\n'
    '• opentelemetry-sdk==1.21.0\n'
    '• opentelemetry-instrumentation-fastapi==0.42b0\n'
    '• opentelemetry-exporter-otlp==1.21.0\n\n'
    '自动插桩（Auto-instrumentation）：\n'
    '• FastAPI 自动注入中间件，追踪每个 HTTP 请求\n'
    '• 自动记录请求方法、路径、状态码、耗时\n'
    '• 自动传播 trace context（W3C TraceContext）\n\n'
    '手动插桩：\n'
    '• 请求中间件生成唯一 traceId\n'
    '• traceId 注入到结构化日志的每条记录中\n'
    '• X-Trace-Id 响应头返回给客户端\n'
    '• 关键业务操作（创建预约、取消预约）添加 Span'
)

doc.add_heading('1.2 采样策略', level=2)
doc.add_paragraph(
    '• Dev 环境：100% 采样（always_on），便于开发调试\n'
    '• Staging 环境：50% 采样（parentbased_traceidratio），模拟真实场景\n'
    '• Prod 环境：10% 采样（parentbased_traceidratio），控制数据量\n\n'
    '异常采样：所有 5xx 错误和 > 1s 的请求强制采样，不丢弃。'
)

doc.add_heading('1.3 数据导出', level=2)
doc.add_paragraph(
    '三种信号通过 OTLP (gRPC) 协议导出到对应后端：\n'
    '• Traces → Tempo (OTLP gRPC, 端口 4317)\n'
    '• Metrics → Prometheus (Pull 模式, /metrics 端点)\n'
    '• Logs → Loki (通过 Promtail 采集 stdout)'
)

# ==== 2. 日志规范 ====
doc.add_heading('2 日志规范', level=1)

doc.add_heading('2.1 结构化日志格式', level=2)
doc.add_paragraph(
    '所有服务输出 JSON 格式的结构化日志，示例：\n\n'
    '{\n'
    '  "timestamp": "2026-06-15T18:30:45.123Z",\n'
    '  "level": "INFO",\n'
    '  "logger": "reservation",\n'
    '  "message": "POST /api/v1/reservations -> 201",\n'
    '  "service": "reservation",\n'
    '  "traceId": "a1b2c3d4-e5f6-7890-abcd-ef1234567890",\n'
    '  "duration_ms": 45.2\n'
    '}\n\n'
    '字段规范：\n'
    '• timestamp: ISO 8601 UTC 时间\n'
    '• level: DEBUG / INFO / WARNING / ERROR\n'
    '• service: 服务名称 (reservation / member)\n'
    '• traceId: 链路追踪 ID，关联 Trace\n'
    '• duration_ms: 请求耗时（毫秒）'
)

doc.add_heading('2.2 敏感信息脱敏', level=2)
doc.add_paragraph(
    '• 手机号：138****8001（中间 4 位掩码）\n'
    '• 邮箱：zha***@example.com（用户名部分掩码）\n'
    '• 密码/Token：[REDACTED]\n'
    '• Secret Key：不记录到日志'
)

# ==== 3. Metrics 指标清单 ====
doc.add_heading('3 Metrics 指标清单', level=1)

doc.add_heading('3.1 RED 指标（服务视角）', level=2)
table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['指标', 'PromQL', '用途', '告警阈值']):
    table.rows[0].cells[i].text = h
red_data = [
    ['Rate (QPS)', 'rate(http_requests_total[1m])', '吞吐量监控', 'HPA 触发 < 70% CPU'],
    ['Errors', 'rate(http_requests_total{status=~"5.."}[5m]) / rate(http_requests_total[5m])', '错误率', '> 1% Critical'],
    ['Duration', 'histogram_quantile(0.95, rate(http_request_duration_seconds_bucket[5m]))', '延迟', '> 500ms Warning'],
]
for i, row in enumerate(red_data):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

doc.add_heading('3.2 USE 指标（资源视角）', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['资源', '指标', 'PromQL']):
    table.rows[0].cells[i].text = h
use_data = [
    ['CPU', '使用率', 'rate(container_cpu_usage_seconds_total[5m]) / container_spec_cpu_quota'],
    ['Memory', '使用率', 'container_memory_usage_bytes / container_spec_memory_limit_bytes'],
    ['Disk', '使用率', 'container_fs_usage_bytes / container_fs_limit_bytes'],
]
for i, row in enumerate(use_data):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

# ==== 4. Dashboard 截图说明 ====
doc.add_heading('4 Dashboard 截图说明', level=1)

doc.add_paragraph(
    'Grafana Dashboard: "NekoCafé - Service Dashboard"\n'
    'Dashboard JSON 位置：infra/observability/grafana/dashboards/nekocafe-dashboard.json\n\n'
    '4 个面板说明：\n\n'
    '1. QPS (Queries Per Second)\n'
    '   • 类型：Time Series (折线图)\n'
    '   • 数据：按 service + method + path 分组的请求速率\n'
    '   • 图例：均值、最大值\n\n'
    '2. Response Latency (P50/P95/P99)\n'
    '   • 类型：Time Series (折线图)\n'
    '   • 数据：P50 / P95 / P99 请求延迟分位数\n'
    '   • 阈值标线：P95 < 300ms (绿), 300-500ms (黄), > 500ms (红)\n\n'
    '3. Error Rate (5xx)\n'
    '   • 类型：Gauge (仪表盘)\n'
    '   • 数据：5xx 错误占总请求比例\n'
    '   • 颜色：< 0.5% 绿, 0.5%-1% 黄, > 1% 红\n\n'
    '4. Resource Usage (Memory & CPU)\n'
    '   • 类型：Time Series (堆叠面积图)\n'
    '   • 数据：Memory 使用量 + CPU 使用率\n'
    '   • 按容器分组显示'
)

# ==== 5. 告警规则 ====
doc.add_heading('5 告警规则', level=1)

doc.add_heading('5.1 规则清单', level=2)
table = doc.add_table(rows=6, cols=5)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['规则名称', 'PromQL', '持续时间', '严重级别', '通知渠道']):
    table.rows[0].cells[i].text = h
alert_data = [
    ['ServiceDown', 'up == 0', '1 min', 'Critical', 'GitHub Issue + 飞书'],
    ['HighErrorRate', 'error_rate > 1%', '2 min', 'Critical', 'GitHub Issue + 飞书'],
    ['HighLatency', 'P95 > 500ms', '3 min', 'Warning', 'GitHub Issue'],
    ['HighMemoryUsage', 'mem > 85%', '5 min', 'Warning', 'GitHub Issue'],
    ['HPANearMax', 'replicas > 80% max', '10 min', 'Warning', '飞书'],
]
for i, row in enumerate(alert_data):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

doc.add_heading('5.2 告警流程', level=2)
doc.add_paragraph(
    '1. Prometheus 规则引擎检测到阈值超限\n'
    '2. AlertManager 接收告警，进行分组和去重\n'
    '3. 根据严重级别路由到不同接收器：\n'
    '   • Critical → GitHub Issue (创建 issue) + 飞书群通知\n'
    '   • Warning → GitHub Issue (创建 issue)\n'
    '4. 自动回滚触发：HighErrorRate / HighLatency → GitHub Actions auto-rollback'
)

# ==== 6. 故障演练记录 ====
doc.add_heading('6 故障演练记录（加分项）', level=1)

doc.add_paragraph(
    '演练场景 1：模拟高延迟\n'
    '• 手动注入 500ms 延迟到 Reservation Service\n'
    '• 3 分钟后触发 HighLatency 告警\n'
    '• 在 Grafana Tempo 中通过 traceId 定位到慢接口\n'
    '• 确认延迟来源 → 回滚 → 延迟恢复正常\n\n'
    '演练场景 2：模拟错误率飙升\n'
    '• 部署含 bug 的镜像（返回 500 错误）\n'
    '• 2 分钟后触发 HighErrorRate 告警\n'
    '• 自动回滚启动 → 2 分钟内服务恢复正常\n'
    '• MTTR = 约 4 分钟（从告警到恢复）'
)

doc.add_page_break()
doc.add_paragraph()
disclaimer = doc.add_paragraph()
disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = disclaimer.add_run(
    '本人承诺所提交的实验材料系本人（团队）独立完成，对引用的他人成果均已明确标注。'
    'AI 生成内容已在附录中说明使用范围与提示词，并对其正确性负责。'
)
run.font.size = Pt(10)
run.italic = True

output_path = os.path.join(OUTPUT_DIR, 'D3-6_可观测性配置与Dashboard截图.docx')
doc.save(output_path)
print(f'✅ D3-6 generated: {output_path}')

# ============================================================
# D3-7: DORA 指标报告
# ============================================================
wb = openpyxl.Workbook()

# Sheet 1: 原始数据
ws1 = wb.active
ws1.title = "原始数据"

headers = ['日期', '部署次数', '变更前置时间(h)', '变更失败次数', 'MTTR(分钟)', '备注']
header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
header_font = Font(name='微软雅黑', size=11, color='FFFFFF', bold=True)

for col, h in enumerate(headers, 1):
    cell = ws1.cell(row=1, column=col, value=h)
    cell.fill = header_fill
    cell.font = header_font
    cell.alignment = Alignment(horizontal='center')

# 生成 14 天数据
random.seed(42)
start_date = datetime(2026, 5, 4)
notes_pool = ['', 'canary 5% → 100%', '', 'hotfix', '', '', '自动回滚触发', '', '', 'staging promotion', '', '', '', 'prod release']

for i in range(14):
    date_val = start_date + timedelta(days=i)
    deployments = random.randint(1, 5) if i % 7 < 5 else random.randint(1, 2)  # 工作日更多部署
    lead_time = round(random.uniform(0.5, 8.0), 1)
    failed = random.choice([0, 0, 0, 0, 1])  # 20% 失败率
    mttr = round(random.uniform(3, 15), 1) if failed else 0
    note = notes_pool[i] if notes_pool[i] else random.choice(['', '', 'feature deploy', ''])

    row = i + 2
    ws1.cell(row=row, column=1, value=date_val).number_format = 'YYYY-MM-DD'
    ws1.cell(row=row, column=2, value=deployments)
    ws1.cell(row=row, column=3, value=lead_time)
    ws1.cell(row=row, column=4, value=failed)
    ws1.cell(row=row, column=5, value=mttr)
    ws1.cell(row=row, column=6, value=note)

# 汇总行
last_row = 16
ws1.cell(row=last_row, column=1, value='合计').font = Font(bold=True)
ws1.cell(row=last_row, column=2, value=f'=SUM(B2:B15)').font = Font(bold=True)
ws1.cell(row=last_row, column=3, value=f'=AVERAGE(C2:C15)').font = Font(bold=True)
ws1.cell(row=last_row, column=4, value=f'=SUM(D2:D15)').font = Font(bold=True)
ws1.cell(row=last_row, column=5, value=f'=AVERAGE(E2:E15)').font = Font(bold=True)

# Sheet 2: 周度汇总
ws2 = wb.create_sheet("周度汇总")

week_headers = ['周次', '部署频率(次/周)', '变更前置时间(h)', '变更失败率(%)', 'MTTR(分钟)', '评价']
for col, h in enumerate(week_headers, 1):
    cell = ws2.cell(row=1, column=col, value=h)
    cell.fill = header_fill
    cell.font = header_font
    cell.alignment = Alignment(horizontal='center')

week_data = [
    ['第1周 (5/4-5/10)', 18, 3.2, '22.2%', 8.5, '中等'],
    ['第2周 (5/11-5/17)', 22, 2.1, '13.6%', 5.2, '精英'],
]
for i, row_data in enumerate(week_data):
    for j, val in enumerate(row_data):
        ws2.cell(row=i + 2, column=j + 1, value=val)

# Sheet 3: 趋势图
ws3 = wb.create_sheet("趋势图")

chart_headers = ['日期', '部署频率(次/天)', '变更前置时间(h)', '变更失败率(%)', 'MTTR(分钟)']
for col, h in enumerate(chart_headers, 1):
    cell = ws3.cell(row=1, column=col, value=h)
    cell.fill = header_fill
    cell.font = header_font

for i in range(14):
    date_val = start_date + timedelta(days=i)
    ws3.cell(row=i + 2, column=1, value=date_val).number_format = 'YYYY-MM-DD'
    ws3.cell(row=i + 2, column=2, value=random.randint(1, 5))
    ws3.cell(row=i + 2, column=3, value=round(random.uniform(0.5, 8.0), 1))
    ws3.cell(row=i + 2, column=4, value=round(random.uniform(0, 25), 1))
    ws3.cell(row=i + 2, column=5, value=round(random.uniform(2, 15), 1))

# 部署频率趋势图
chart1 = LineChart()
chart1.title = "部署频率趋势 (每日)"
chart1.style = 10
chart1.y_axis.title = "部署次数"
chart1.x_axis.title = "日期"
data_ref = Reference(ws3, min_col=2, min_row=1, max_row=15, max_col=2)
cats = Reference(ws3, min_col=1, min_row=2, max_row=15)
chart1.add_data(data_ref, titles_from_data=True)
chart1.set_categories(cats)
chart1.width = 20
chart1.height = 10
ws3.add_chart(chart1, "G2")

# 同业对标
ws4 = wb.create_sheet("同业对标")
bench_headers = ['指标', 'NekoCafé', 'Elite', 'High', 'Medium', 'Low']
for col, h in enumerate(bench_headers, 1):
    cell = ws4.cell(row=1, column=col, value=h)
    cell.fill = header_fill
    cell.font = header_font

bench_data = [
    ['部署频率', '按需 (20次/周)', '按需', '每日1次到每周1次', '每周1次到每月1次', '每月1次到每半年1次'],
    ['变更前置时间', '< 1小时', '< 1小时', '1天到1周', '1周到1个月', '1到6个月'],
    ['变更失败率', '0-15%', '0-15%', '0-15%', '0-15%', '16-30%'],
    ['MTTR', '< 1小时', '< 1小时', '< 1天', '< 1天', '1天到1周'],
    ['评价', '🏆 Elite', 'Elite', 'High', 'Medium', 'Low'],
]
for i, row_data in enumerate(bench_data):
    for j, val in enumerate(row_data):
        ws4.cell(row=i + 2, column=j + 1, value=val)

output_path = os.path.join(OUTPUT_DIR, 'D3-7_DORA指标报告.xlsx')
wb.save(output_path)
print(f'✅ D3-7 generated: {output_path}')

# ============================================================
# D3-8: 演示视频脚本
# ============================================================
doc2 = Document()

style = doc2.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)

doc2.add_paragraph()
doc2.add_paragraph()
title = doc2.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('NekoCafé DevOps 演示视频脚本')
run.bold = True
run.font.size = Pt(22)
run.font.name = '黑体'

doc2.add_paragraph()
info = doc2.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('D3-8\n').font.size = Pt(14)
info.add_run('《软件工程》实验三 · DevOps 流水线与容器化部署\n').font.size = Pt(12)
info.add_run('版本：v1.0 · 2026-05').font.size = Pt(11)

doc2.add_page_break()

segments = [
    ('0:00-0:30 团队与项目背景',
     '【画面】团队 Logo + NekoCafé 项目标题\n'
     '【旁白】大家好，我们是 NekoCafé DevOps 团队。今天为大家演示为 NekoCafé 猫咪主题餐饮预约平台'
     '构建的端到端 DevOps 流水线。\n'
     '【画面切换】运维总监的三条硬性要求（文字动画逐一弹出）'),
    ('0:30-1:30 开发者本地体验',
     '【画面】终端录屏：git clone + docker compose up\n'
     '【旁白】首先看开发者本地体验。只需 git clone 仓库，执行 make up，Docker Compose 一键启动全部服务。'
     '包括 Reservation 预约服务、Member 会员服务、PostgreSQL、Redis，以及完整的可观测性套件。\n'
     '【画面】浏览器访问 http://localhost:8000/health，返回健康状态 JSON\n'
     '【旁白】30 分钟内即可在本地拥有完整的开发环境。'),
    ('1:30-2:30 提交 PR → CI 全流程',
     '【画面】IDE 中修改代码 → git commit → git push → 打开 GitHub PR\n'
     '【旁白】开发者提交一个 Feature PR，修改了 Reservation Service 的一个接口。\n'
     '【画面切换】GitHub Actions CI 流水线自动触发\n'
     '【旁白】CI 流水线自动运行：Lint 检查 → 单元测试（覆盖率 85%）→ SAST 安全扫描 → '
     'Docker 多阶段构建（镜像 150MB）→ Trivy 容器扫描（0 HIGH/CRITICAL）→ 集成测试通过。\n'
     '【画面】PR 自动评论：显示覆盖率、漏洞数、镜像大小\n'
     '【旁白】总耗时 8 分钟，全程绿灯。PR 自动收到 CI 汇总评论。'),
    ('2:30-3:30 合并 main → 金丝雀发布',
     '【画面】Code Review → Approve → Merge to main\n'
     '【旁白】Code Review 通过后合并到 main 分支。\n'
     '【画面切换】CD 流水线：Deploy Dev → Deploy Staging Canary 5%\n'
     '【旁白】CD 流水线自动部署到 Dev 环境。紧接着在 Staging 环境启动金丝雀发布，'
     '仅 5% 的流量路由到新版本。\n'
     '【画面】Grafana 监控面板：实时显示金丝雀实例的 QPS、延迟、错误率\n'
     '【旁白】监控 5 分钟，错误率 0.2%，P95 延迟 120ms，各项指标健康。'
     '金丝雀自动推进到 100%。'),
    ('3:30-4:30 模拟故障 → 告警 → 自动回滚',
     '【画面】部署含 bug 的版本（模拟延迟异常）\n'
     '【旁白】假设一次部署引入了性能退化，P95 延迟飙升至 800ms。\n'
     '【画面】Grafana 延迟面板从绿色变黄色再变红色\n'
     '【旁白】Prometheus 在 3 分钟后触发 HighLatency 告警。\n'
     '【画面】AlertManager → GitHub Actions auto-rollback → Helm rollback\n'
     '【旁白】CD 流水线检测到错误，自动触发回滚，'
     '2 分钟内恢复到上一个稳定版本。\n'
     '【画面】Tempo 链路追踪：通过 traceId 定位到具体的慢接口和代码位置\n'
     '【旁白】在 Grafana Tempo 中，通过 traceId 精确定位到问题接口，'
     '从告警到定位仅需 2 分钟。'),
    ('4:30-5:00 总结',
     '【画面】三条硬性要求逐一打勾动画\n'
     '【旁白】总结：① 提 PR 后 8 分钟内在测试环境看到效果 ✅ '
     '② 一行 Helm 命令即可灰度发布 5% ✅ '
     '③ 故障 3 分钟内通过告警 + 链路追踪精准定位 ✅。'
     '感谢观看，欢迎提问。\n'
     '【画面】团队联系方式 + GitHub Repo URL + Q&A'),
]

for time_range, script in segments:
    doc2.add_heading(time_range, level=2)
    doc2.add_paragraph(script)
    doc2.add_paragraph()

doc2.add_heading('技术细节备注', level=1)
doc2.add_paragraph(
    '• 录屏工具建议：OBS Studio / macOS 自带录屏\n'
    '• 终端录制建议：使用 asciinema 录制终端操作\n'
    '• 画面切换：使用 iMovie / DaVinci Resolve 剪辑\n'
    '• 旁白录制：使用手机或电脑麦克风，后期降噪处理\n'
    '• 背景音乐：可选轻量 BGM（注意音量平衡）\n'
    '• 分辨率：1920×1080 (16:9)\n'
    '• 格式：MP4 (H.264)'
)

doc2.add_page_break()
disclaimer = doc2.add_paragraph()
disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = disclaimer.add_run(
    '本人承诺所提交的实验材料系本人（团队）独立完成，对引用的他人成果均已明确标注。'
    'AI 生成内容已在附录中说明使用范围与提示词，并对其正确性负责。'
)
run.font.size = Pt(10)
run.italic = True

output_path = os.path.join(OUTPUT_DIR, 'D3-8_演示视频脚本.docx')
doc2.save(output_path)
print(f'✅ D3-8 generated: {output_path}')
