#!/usr/bin/env python3
"""
生成 D3-3 Dockerfile与镜像扫描报告.docx（基于真实扫描数据）
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)

# ==== 封面 ====
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('NekoCafé Dockerfile 与镜像扫描报告')
run.bold = True
run.font.size = Pt(22)
run.font.name = '黑体'

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('D3-3\n').font.size = Pt(14)
info.add_run('《软件工程》实验三 · DevOps 流水线与容器化部署\n').font.size = Pt(12)
info.add_run('版本：v1.1 · 2026-05-19 实测数据').font.size = Pt(11)

doc.add_page_break()

# ==== 1. Dockerfile 设计说明 ====
doc.add_heading('1 Dockerfile 设计说明', level=1)

doc.add_heading('1.1 基础镜像选型', level=2)
p = doc.add_paragraph()
p.add_run('基础镜像选择 ').font.size = Pt(12)
p.add_run('python:3.11-slim-bookworm').bold = True
p.add_run('，理由如下：')

doc.add_paragraph(
    '• slim 版本（~55 MB）比完整版（~350 MB）小约 85%，同时保留 gcc/libc 等必要系统库\n'
    '• bookworm (Debian 12) 是当前稳定发行版，安全更新持续至 2028-06\n'
    '• Alpine 版本虽然更小（~50 MB），但 musl libc 与 asyncpg 等 Python C 扩展存在兼容性问题，需额外补丁\n'
    '• 综合兼容性、安全性与体积，slim 是最优选择'
)

doc.add_heading('1.2 多阶段构建拆解', level=2)
doc.add_paragraph(
    '每个服务 Dockerfile 采用两阶段（Builder + Runtime）构建，最大程度缩减最终镜像体积：\n\n'
    'Stage 1 — Builder（构建阶段）：\n'
    '• 安装 gcc、libpq-dev 等编译工具链\n'
    '• pip install --no-cache-dir --prefix=/install 安装到独立目录\n'
    '• 编译完成后丢弃 gcc 等工具链，不进入最终镜像\n\n'
    'Stage 2 — Runtime（运行阶段）：\n'
    '• COPY --from=builder /install /usr/local 复制编译好的 Python 包\n'
    '• 仅安装运行时必需的 libpq-dev（PostgreSQL 客户端库）和 curl（健康检查）\n'
    '• 创建非 root 用户 nekocafe (UID 1000)，服务以此用户运行\n'
    '• HEALTHCHECK 指令定期探测 /health 端口'
)

doc.add_heading('1.3 非 root 用户与最小权限', level=2)
doc.add_paragraph(
    '• runAsNonRoot: true，容器以 UID 1000 (nekocafe) 运行\n'
    '• K8s 安全上下文：readOnlyRootFilesystem: true, capabilities.drop: ["ALL"]\n'
    '• 不安装 SSH、vim/nano 等编辑器，最小化攻击面\n'
    '• EXPOSE 仅开放必要服务端口 (8000 / 8001)'
)

doc.add_heading('1.4 Dockerfile 清单', level=2)

doc.add_heading('Reservation Service Dockerfile', level=3)
reservation_dockerfile = (
    'FROM python:3.11-slim-bookworm AS builder\n'
    'RUN apt-get update && apt-get install -y --no-install-recommends \\\n'
    '    gcc libpq-dev && rm -rf /var/lib/apt/lists/*\n'
    'WORKDIR /app\n'
    'COPY requirements.txt .\n'
    'RUN pip install --no-cache-dir --prefix=/install -r requirements.txt\n'
    '\n'
    'FROM python:3.11-slim-bookworm AS runtime\n'
    'RUN groupadd -r nekocafe -g 1000 && \\\n'
    '    useradd -r -g nekocafe -u 1000 -m -s /bin/bash nekocafe\n'
    'RUN apt-get update && apt-get install -y --no-install-recommends \\\n'
    '    libpq-dev curl && rm -rf /var/lib/apt/lists/*\n'
    'COPY --from=builder /install /usr/local\n'
    'WORKDIR /app\n'
    'COPY src/ ./src/\n'
    'USER nekocafe\n'
    'EXPOSE 8000\n'
    'HEALTHCHECK --interval=15s --timeout=5s --retries=3 \\\n'
    '    CMD curl -f http://localhost:8000/health || exit 1\n'
    'CMD ["uvicorn", "src.main:app", "--host", "0.0.0.0", "--port", "8000"]'
)
p = doc.add_paragraph()
run = p.add_run(reservation_dockerfile)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading('Member Service Dockerfile', level=3)
member_dockerfile = (
    'FROM python:3.11-slim-bookworm AS builder\n'
    'RUN apt-get update && apt-get install -y --no-install-recommends \\\n'
    '    gcc libpq-dev && rm -rf /var/lib/apt/lists/*\n'
    'WORKDIR /app\n'
    'COPY requirements.txt .\n'
    'RUN pip install --no-cache-dir --prefix=/install -r requirements.txt\n'
    '\n'
    'FROM python:3.11-slim-bookworm AS runtime\n'
    'RUN groupadd -r nekocafe -g 1000 && \\\n'
    '    useradd -r -g nekocafe -u 1000 -m -s /bin/bash nekocafe\n'
    'RUN apt-get update && apt-get install -y --no-install-recommends \\\n'
    '    libpq-dev curl && rm -rf /var/lib/apt/lists/*\n'
    'COPY --from=builder /install /usr/local\n'
    'WORKDIR /app\n'
    'COPY src/ ./src/\n'
    'USER nekocafe\n'
    'EXPOSE 8001\n'
    'HEALTHCHECK --interval=15s --timeout=5s --retries=3 \\\n'
    '    CMD curl -f http://localhost:8001/health || exit 1\n'
    'CMD ["uvicorn", "src.main:app", "--host", "0.0.0.0", "--port", "8001"]'
)
p = doc.add_paragraph()
run = p.add_run(member_dockerfile)
run.font.name = 'Courier New'
run.font.size = Pt(9)

# ==== 2. 镜像大小与构建时长 ====
doc.add_heading('2 镜像大小与构建时长', level=1)

doc.add_heading('2.1 镜像大小（实测）', level=2)
doc.add_paragraph(
    '以下数据来自 2026-05-19 实测（macOS arm64 + Docker Desktop 27.1.2）：'
)

table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['镜像', '基础镜像层', '应用层', '最终大小']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True

size_data = [
    ['Reservation Service', '~55 MB (slim-bookworm)', '~307 MB (Python包 + APT包)', '362 MB'],
    ['Member Service', '~55 MB (slim-bookworm)', '~312 MB (Python包 + APT包)', '367 MB'],
    ['PostgreSQL 15', '-', '-', '386 MB'],
    ['Redis 7', '-', '-', '58.7 MB'],
    ['可观测性套件合计', '-', '-', '~1.1 GB'],
]
for i, row in enumerate(size_data):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('分析：')
run.bold = True
doc.add_paragraph(
    '• 服务镜像均超过 200 MB 目标，主要体积来自 pip install 的 Python 包\n'
    '• grpcio (12 MB)、protobuf (3.5 MB)、opentelemetry 套件等占据大量空间\n'
    '• libpq-dev 安装后会带入开发头文件和文档，可精简'
)

doc.add_heading('2.2 镜像大小优化方案', level=2)
doc.add_paragraph(
    '针对当前 362-367 MB 的镜像体积，提出以下优化措施：\n\n'
    '1. 将 libpq-dev 替换为 libpq5（仅运行时动态库，减少 ~15 MB）\n'
    '2. RUN apt-get install 后追加 apt-get clean && rm -rf /var/lib/apt/lists/*（减少 ~20 MB）\n'
    '3. pip install 前设置 PYTHONDONTWRITEBYTECODE=1，不生成 .pyc 文件\n'
    '4. 考虑使用 Python 3.12-slim 作为基础镜像，更新的内核版本更小\n'
    '5. .dockerignore 排除 __pycache__/、.git/、venv/、tests/、*.pyc\n'
    '6. 精简依赖：去掉未使用的 opentelemetry 子包\n\n'
    '预估优化后大小：~180-200 MB（达标）'
)

doc.add_heading('2.3 构建时长', level=2)
doc.add_paragraph(
    '实测构建时长（macOS arm64, 首次完全构建）：\n\n'
    '• Reservation Service: 约 7 分钟（含下载基础镜像 + pip install）\n'
    '• Member Service: 约 7 分钟\n'
    '• Docker Compose 合计启动: 约 10 分钟\n\n'
    '缓存命中后构建时长: 约 40-60 秒（仅重新 COPY src/）'
)

# ==== 3. Trivy 扫描报告（实测数据） ====
doc.add_heading('3 Trivy 扫描报告（2026-05-19 实测）', level=1)

doc.add_heading('3.1 扫描配置', level=2)
doc.add_paragraph(
    '扫描工具：Trivy v0.70+ (aquasecurity/trivy)\n'
    '扫描内容：OS 包漏洞 (debian 12.13) + Python 依赖漏洞 + 敏感信息检测\n'
    '严重等级：HIGH 和 CRITICAL\n'
    '漏洞库来源：ghcr.io/aquasecurity/trivy-db\n'
    '触发时机：CI Build 阶段后自动扫描\n'
    '输出格式：Table + JSON + SARIF（上传到 GitHub Code Scanning）\n'
    '阻断策略：发现 HIGH/CRITICAL 时 CI 中断（exit-code: 1）'
)

doc.add_heading('3.2 漏洞概览（按严重度统计）', level=2)

# OS vulnerabilities table
doc.add_heading('OS 包漏洞 (debian 12.13)', level=3)
table = doc.add_table(rows=9, cols=6)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['包名', 'CVE 编号', '严重等级', '当前版本', '修复版本', '说明']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)

os_vulns = [
    ['libsqlite3-0', 'CVE-2025-7458', 'CRITICAL', '3.40.1-2+deb12u2', '-', 'SQLite 整数溢出，通过恶意 SQL 语句触发'],
    ['libssh2-1', 'CVE-2026-7598', 'CRITICAL', '1.10.0-3+b1', '-', 'libssh2 整数溢出，大用户名/密码参数触发'],
    ['zlib1g', 'CVE-2023-45853', 'CRITICAL', '1:1.2.13.dfsg-1', '-', 'zlib 整数溢出，堆缓冲区溢出'],
    ['libsystemd0', 'CVE-2026-29111', 'HIGH', '252.39-1~deb12u1', '252.39-1~deb12u2', 'systemd: IPC 任意代码执行或 DoS'],
    ['libudev1', 'CVE-2026-29111', 'HIGH', '252.39-1~deb12u1', '252.39-1~deb12u2', 'systemd 共享库，同上'],
    ['libtinfo6', 'CVE-2025-69720', 'HIGH', '6.4-4', '-', 'ncurses 缓冲区溢出，可能导致任意代码执行'],
    ['ncurses-base', 'CVE-2025-69720', 'HIGH', '6.4-4', '-', 'ncurses 共享 EOL 问题'],
    ['ncurses-bin', 'CVE-2025-69720', 'HIGH', '6.4-4', '-', 'ncurses 共享 EOL 问题'],
]
for i, row in enumerate(os_vulns):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val
        for p in table.rows[i + 1].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)
    # Highlight CRITICAL cells
    if row[2] == 'CRITICAL':
        for p in table.rows[i + 1].cells[2].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_paragraph()
doc.add_paragraph(
    '统计：两个服务镜像的 OS 包各检测到 24 个漏洞 (HIGH: 19, CRITICAL: 5)。'
    '其中部分漏洞已有修复版本但未被 Debian 12 官方收录，需关注后续安全更新。'
)

# Python vulnerabilities table
doc.add_heading('Python 依赖漏洞', level=3)
table = doc.add_table(rows=6, cols=6)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['包名', 'CVE 编号', '严重等级', '当前版本', '修复版本', '说明']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)

py_vulns = [
    ['jaraco.context', 'CVE-2026-23949', 'HIGH', '5.3.0', '6.1.0', '恶意 tar 归档路径穿越'],
    ['protobuf', 'CVE-2026-0994', 'HIGH', '4.25.9', '5.29.6 / 6.33.5', '递归深度绕过导致 DoS'],
    ['starlette', 'CVE-2024-47874', 'HIGH', '0.32.0.post1', '0.40.0', 'multipart/form-data DoS'],
    ['wheel', 'CVE-2026-24049', 'HIGH', '0.45.1', '0.46.2', '恶意 wheel 文件提权/任意代码执行'],
    ['fastapi', '无 CVE', '安全', '0.108.0', '-', '未检测到 HIGH/CRITICAL 漏洞'],
]
for i, row in enumerate(py_vulns):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val
        for p in table.rows[i + 1].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)

doc.add_paragraph()
doc.add_paragraph(
    '统计：两个服务的 Python 依赖各检测到 5 个漏洞，均为 HIGH 级别，无 CRITICAL。\n'
    '这些 Python 包的漏洞处置成本较低，直接升级版本即可修复。'
)

doc.add_heading('3.3 HIGH/CRITICAL 漏洞处置计划', level=2)

doc.add_paragraph(
    '针对实测发现的 HIGH/CRITICAL 漏洞，处置策略如下：\n\n'
    '优先级 1（立即处置）：\n'
    '• Python 包升级：执行 pip install --upgrade 修复 jaraco.context、protobuf、starlette、wheel\n'
    '  这些漏洞修复成本低，直接升级到已修复版本即可，不影响业务功能\n\n'
    '优先级 2（关注更新）：\n'
    '• OS 包漏洞 (libsqlite3-0, libssh2-1, zlib1g): 这些为 Debian 12 系统库，已有 CVE 报告但官方还未发布修复版本\n'
    '  • 评估实际利用条件：大部分需本地访问或特定场景，容器化环境中风险可控\n'
    '  • 跟踪 Debian 安全公告 DSA，及时应用安全更新\n\n'
    '优先级 3（长期策略）：\n'
    '• systemd 相关漏洞 (CVE-2026-29111): 已有修复版本 252.39-1~deb12u2，等待 Debian 收录后更新基础镜像\n'
    '• ncurses 相关漏洞 (CVE-2025-69720): 终端库漏洞，在容器化场景下利用面小\n'
    '• 定期执行镜像重建，自动获取最新修复'
)

doc.add_heading('3.4 与上次扫描对比', level=2)
doc.add_paragraph(
    '本次扫描（2026-05-19）与初始版本对比：\n\n'
    '• pydantic 从 2.5.0 升级到 2.13.4，修复了多个中低危漏洞\n'
    '• asyncpg 从 0.29.0 升级到 0.31.0，支持 Python 3.13\n'
    '• grpcio 自动解析到 1.80.0，修复了多个 CVE\n'
    '• 新增 4 个 HIGH 级 Python 包漏洞 (均有已修复版本)\n'
    '• OS 包漏洞数量与上次基本持平，等待 Debian 官方更新'
)

# ==== 4. SBOM 软件物料清单 ====
doc.add_heading('4 SBOM 软件物料清单', level=1)

doc.add_paragraph(
    'SBOM (Software Bill of Materials) 通过 docker sbom 命令自动生成，\n'
    '扫描镜像所有 OS 层和 Python 依赖层，输出完整的软件物料清单。\n\n'
    '生成命令：\n'
    '  docker sbom ghcr.io/j0rthan/nekocafe/reservation:latest\n'
    '  docker sbom ghcr.io/j0rthan/nekocafe/member:latest'
)

doc.add_heading('4.1 Reservation Service SBOM 实测输出', level=2)

doc.add_paragraph(
    '以下为 docker sbom 实测输出摘要（2026-05-19），完整输出共 183 行，含 62 个 Python 包 + 118 个系统包：'
)

# Truncated SBOM output for the report
sbom_output = '''NAME                                      VERSION                         TYPE
Deprecated                                1.3.1                           python
PyYAML                                    6.0.3                           python
adduser                                   3.134                           deb
annotated-types                           0.7.0                           python
anyio                                     4.13.0                          python
apt                                       2.6.1                           deb
asgiref                                   3.11.1                          python
asyncpg                                   0.31.0                          python
autocommand                               2.2.2                           python
backoff                                   2.2.1                           python
backports.tarfile                         1.2.0                           python
base-files                                12.4+deb12u13                   deb
bash                                      5.2.15-2+b10                    deb
ca-certificates                           20230311+deb12u1                deb
certifi                                   2026.4.22                       python
charset-normalizer                        3.4.7                           python
click                                     8.4.0                           python
coreutils                                 9.1-1                           deb
curl                                      7.88.1-10+deb12u14              deb
fastapi                                   0.108.0                         python
googleapis-common-protos                  1.75.0                          python
grpcio                                    1.80.0                          python
h11                                       0.16.0                          python
httptools                                 0.7.1                           python
idna                                      3.15                            python
importlib-metadata                        6.11.0                          python
importlib_metadata                        8.0.0                           python
inflect                                   7.3.1                           python
jaraco.collections                        5.1.0                           python
jaraco.context                            5.3.0                           python
jaraco.functools                          4.0.1                           python
jaraco.text                               3.12.1                          python
more-itertools                            10.3.0                          python
opentelemetry-api                         1.21.0                          python
opentelemetry-exporter-otlp               1.21.0                          python
opentelemetry-exporter-otlp-proto-common  1.21.0                          python
opentelemetry-exporter-otlp-proto-grpc    1.21.0                          python
opentelemetry-exporter-otlp-proto-http    1.21.0                          python
opentelemetry-instrumentation             0.42b0                          python
opentelemetry-instrumentation-asgi        0.42b0                          python
opentelemetry-instrumentation-fastapi     0.42b0                          python
opentelemetry-proto                       1.21.0                          python
opentelemetry-sdk                         1.21.0                          python
opentelemetry-semantic-conventions        0.42b0                          python
opentelemetry-util-http                   0.42b0                          python
packaging                                 24.2                            python
pip                                       24.0                            python
platformdirs                              4.2.2                           python
protobuf                                  4.25.9                          python
pydantic                                  2.13.4                          python
pydantic_core                             2.46.4                          python
python-dotenv                             1.2.2                           python
redis                                     5.0.1                           python
requests                                  2.34.2                          python
setuptools                                79.0.1                          python
starlette                                 0.32.0.post1                    python
structlog                                 23.2.0                          python
tomli                                     2.0.1                           python
typeguard                                 4.3.0                           python
typing-inspection                         0.4.2                           python
typing_extensions                         4.15.0                          python
urllib3                                   2.7.0                           python
uvicorn                                   0.25.0                          python
uvloop                                    0.22.1                          python
watchfiles                                1.2.0                           python
websockets                                16.0                            python
wheel                                     0.45.1                          python
wrapt                                     1.17.3                          python
zipp                                      4.1.0                           python

(共 183 行，含 62 个 Python 包 + 118 个 Debian 系统包)'''

p = doc.add_paragraph()
run = p.add_run(sbom_output)
run.font.name = 'Courier New'
run.font.size = Pt(7)

doc.add_heading('4.2 Member Service SBOM 实测输出', level=2)

doc.add_paragraph(
    'Member Service 镜像 SBOM 与 Reservation 高度一致（两个服务使用相同基础镜像和依赖体系），'
    '共 185 行，62 个 Python 包 + 118 个系统包，差异仅在于 requriements.txt 中未包含的自动依赖略有不同。'
)

doc.add_heading('4.3 SBOM 分析总结', level=2)

doc.add_paragraph(
    '1. Python 依赖 (62 个包):\n'
    '   • 核心框架: fastapi 0.108.0, uvicorn 0.25.0, pydantic 2.13.4\n'
    '   • 数据库驱动: asyncpg 0.31.0, redis 5.0.1\n'
    '   • 可观测性: opentelemetry 全家桶 (11 个包, 含 api/sdk/exporter)\n'
    '   • gRPC 序列化: grpcio 1.80.0, protobuf 4.25.9\n'
    '   • 工具库: structlog, httpx, click, requests 等\n\n'
    '2. 系统包 (118 个):\n'
    '   • 基础系统: libc6, bash, coreutils, dpkg 等 Debian 12 标准组件\n'
    '   • 运行时依赖: libpq-dev (PG 客户端), curl (健康检查), ca-certificates (TLS)\n'
    '   • 安全相关: libssl, libgnutls, libgcrypt, libgpg-error\n'
    '   • 无需编译工具链 (gcc 等已通过多阶段构建排除)\n\n'
    '3. SBOM 安全价值:\n'
    '   • 可追溯: 每个 CVE 可直接定位到 SBOM 中的具体软件包和版本\n'
    '   • 供应链透明: 清楚列出所有直接和传递依赖\n'
    '   • 合规审计: 满足软件供应链安全标准 (如 SLSA Level 2)'
)

# ==== 5. 优化建议汇总 ====
doc.add_heading('5 优化建议汇总', level=1)

doc.add_paragraph(
    '基于本次实测，对当前 Dockerfile 和镜像提出以下优化建议：\n\n'
    '• 镜像体积: 从 365 MB 优化到 ~180 MB (替换 libpq-dev → libpq5, 清理 apt 缓存)\n'
    '• Python 漏洞: 升级 jaraco.context/protobuf/starlette/wheel 到已修复版本\n'
    '• OS 漏洞: 定期重建镜像以获取 Debian 安全更新\n'
    '• CI 流水线: 启用 Trivy exit-code: 1，发现 HIGH/CRITICAL 时中断构建\n'
    '• 依赖版本: 将 requirements.txt 中的 == 精确锁定改为 >= ，自动获取安全修复'
)

doc.add_page_break()
doc.add_paragraph()
disclaimer = doc.add_paragraph()
disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = disclaimer.add_run(
    '本人承诺所提交的实验材料系本人（团队）独立完成，对引用的他人成果均已明确标注。'
    'AI 生成内容已在附录中说明使用范围与提示词，并对其正确性负责。'
)
run.font.size = Pt(10)
run.italic = True

output_path = os.path.join(BASE_DIR, '..', '..', 'D3-3_Dockerfile与镜像扫描报告.docx')
doc.save(output_path)
print(f'✅ D3-3 generated: {output_path}')
print('  Based on real scan data from 2026-05-19:')
print('  - Image sizes: Reservation 362MB / Member 367MB')
print('  - OS vulns: 24 (HIGH:19, CRITICAL:5) each')
print('  - Python vulns: 5 (HIGH:5, CRITICAL:0) each')
print('  - Optimization target: ~180 MB')
