#!/usr/bin/env python3
"""
生成 D3-1 DevOps设计方案.docx
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# 页眉页脚
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)

# ==== 封面 ====
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('NekoCafé DevOps 设计方案')
run.bold = True
run.font.size = Pt(26)
run.font.name = '黑体'

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('DevOps Pipeline & Containerized Deployment Design')
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('北京林业大学 · 信息学院\n').font.size = Pt(14)
info.add_run('《软件工程》课程 · 实验三\n').font.size = Pt(14)
info.add_run('版本：v1.0 · 2026-05').font.size = Pt(12)

doc.add_page_break()

# ==== 第1章：概述 ====
doc.add_heading('1 概述', level=1)

doc.add_heading('1.1 项目目标', level=2)
doc.add_paragraph(
    'NekoCafé 是一家面向年轻用户的连锁猫咪主题餐厅集团，目前在全国 12 个城市开设 38 家门店。'
    '本次 DevOps 方案的目标是为 NekoCafé 新一代智慧餐饮预约平台构建端到端的 CI/CD 流水线，'
    '实现从代码提交到生产部署的全自动化，支撑运维总监提出的三条硬性要求：'
    '① 提 PR 后 10 分钟内在测试环境看到效果；'
    '② 一行配置即可灰度发布到 5% 的门店；'
    '③ 出问题 3 分钟内能定位到具体服务、具体节点、具体接口。'
)

doc.add_heading('1.2 范围界定', level=2)
doc.add_paragraph(
    '本 DevOps 方案覆盖以下范围：\n'
    '• 选取「预约服务」+「会员服务」两个核心服务作为 PoC\n'
    '• 覆盖 CI/CD 完整流水线（Lint → Test → SAST → Build → Scan → Integration → Deploy）\n'
    '• 实现金丝雀发布（Canary Release）与自动回滚\n'
    '• 建立可观测性三件套：日志（Loki）+ 指标（Prometheus）+ 链路追踪（Tempo）\n'
    '• 产出 DORA 四指标数据'
)

doc.add_heading('1.3 与实验二架构的衔接', level=2)
doc.add_paragraph(
    '本实验基于实验二的微服务架构设计，选取「预约服务 (Reservation Service)」和「会员服务 (Member Service)」'
    '两个核心微服务进行容器化与 CI/CD 落地。两个服务均采用 FastAPI 框架，通过 RESTful API 通信，'
    '共享 PostgreSQL 数据库（逻辑隔离），使用 Redis 做缓存层。'
)

# ==== 第2章：CALMS 原则 ====
doc.add_heading('2 CALMS 原则在本项目中的落地', level=1)

calms_data = [
    ('Culture（文化）',
     '• 所有配置、脚本、文档均纳入 Git 版本管理\n'
     '• PR 强制 Code Review，至少一位 Approver\n'
     '• 开发与运维共享 Dashboard，透明的监控数据\n'
     '• 建立 Blameless Postmortem 文化'),
    ('Automation（自动化）',
     '• CI 流水线全自动：Lint → Test → SAST → Build → Scan → Integration → Push\n'
     '• CD 流水线：自动部署 Dev/Staging，手动审批后部署 Prod\n'
     '• 金丝雀发布自动渐进推进（5% → 25% → 50% → 100%）\n'
     '• 异常自动回滚（P95 > 500ms 或错误率 > 1%）'),
    ('Lean（精益）',
     '• 多阶段 Dockerfile 最小化镜像（≤ 200 MB）\n'
     '• CI 总时长 ≤ 10 分钟\n'
     '• 缓存策略（pip cache, Docker layer cache, GHA cache）\n'
     '• Value Stream Mapping 消除等待'),
    ('Measurement（度量）',
     '• DORA 四指标：部署频率 / 变更前置时间 / 变更失败率 / MTTR\n'
     '• RED 指标：Rate / Errors / Duration\n'
     '• USE 指标：Utilization / Saturation / Errors\n'
     '• Grafana Dashboard 实时展示'),
    ('Sharing（分享）',
     '• Monorepo 架构，代码与配置共享\n'
     '• OpenTelemetry 统一可观测性标准\n'
     '• 运维手册（Runbook）与回滚手册（Rollback）文档化\n'
     '• 答辩 PPT 总结踩坑经验')
]

for title, content in calms_data:
    doc.add_heading(f'2.{calms_data.index((title, content)) + 1} {title}', level=2)
    doc.add_paragraph(content)

# ==== 第3章：分支策略 ====
doc.add_heading('3 分支策略', level=1)

doc.add_heading('3.1 Trunk-Based Development', level=2)
doc.add_paragraph(
    '本项目采用 Trunk-Based Development（基于主干的开发模式）：\n\n'
    '• main 分支：始终保持可部署状态，CI 通过后自动部署 Dev/Staging\n'
    '• feature/* 分支：短期特性分支（< 1 天），完成后立即合并 main\n'
    '• hotfix/* 分支：紧急修复，从 main 分出，修复后双向合并\n\n'
    '选择理由：团队规模小（< 5 人），服务少（2 个），Trunk-Based 比 GitFlow 更高效，'
    '减少合并冲突，加速反馈循环。'
)

doc.add_heading('3.2 分支保护规则', level=2)
doc.add_paragraph(
    '• 禁止直接 push 到 main 分支\n'
    '• 合并前必须 CI 全部通过\n'
    '• 合并前至少 1 人 Code Review 通过\n'
    '• 合并前必须与 main 分支同步（no conflicts）\n'
    '• 禁止绕过 Branch Protection 规则'
)

# ==== 第4章：环境拓扑 ====
doc.add_heading('4 环境拓扑', level=1)

doc.add_heading('4.1 三环境架构', level=2)

table = doc.add_table(rows=5, cols=5)
table.style = 'Light Grid Accent 1'
headers = ['环境', '用途', 'K8s Namespace', '触发方式', '审批']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

envs = [
    ['dev', '开发自测', 'nekocafe-dev', 'PR 自动', '无'],
    ['staging', '预发布验证', 'nekocafe-staging', '合并 main 自动', '无'],
    ['prod (canary)', '金丝雀 5%', 'nekocafe-prod', '手动触发', '需要'],
    ['prod (stable)', '生产全量', 'nekocafe-prod', '金丝雀通过后', '自动'],
]
for i, env in enumerate(envs):
    for j, val in enumerate(env):
        table.rows[i + 1].cells[j].text = val

doc.add_heading('4.2 数据隔离', level=2)
doc.add_paragraph(
    '• Dev 环境：使用本地 Docker Compose 或 Minikube，PostgreSQL 不持久化\n'
    '• Staging 环境：独立的 PostgreSQL 实例，定期从 Prod 脱敏同步\n'
    '• Prod 环境：独立的 PostgreSQL 集群，主从架构，每日自动备份\n'
    '• Secret 管理：通过 GitHub Secrets + K8s Secret 注入，严禁硬编码'
)

# ==== 第5章：流水线总体设计 ====
doc.add_heading('5 流水线总体设计', level=1)

doc.add_heading('5.1 CI 流水线', level=2)
doc.add_paragraph(
    'CI 流水线包含以下阶段，任一失败即中断：\n\n'
    'Lint (flake8 + black + hadolint)\n'
    '  → Unit Test + Coverage (pytest, 覆盖率 ≥ 80%)\n'
    '  → SAST (Bandit 安全扫描)\n'
    '  → Build (Docker 多阶段构建, 镜像 ≤ 200 MB)\n'
    '  → Container Scan (Trivy, 无 HIGH/CRITICAL 漏洞)\n'
    '  → Integration Test (Docker Compose 起栈验证)\n'
    '  → Push Image (推送到 GHCR)\n'
    '  → PR Comment (自动评论覆盖率/漏洞/镜像大小)'
)

doc.add_heading('5.2 CD 流水线', level=2)
doc.add_paragraph(
    'CD 流水线分三阶段：\n\n'
    '1. Deploy to Dev：PR 合并到 main 后自动部署，Helm 更新\n'
    '2. Deploy to Staging (Canary 5%)：金丝雀部署 5% 流量，监控 5 分钟\n'
    '   → 错误率 < 1% 且 P95 < 500ms → 逐步推进\n'
    '   → 超过阈值 → 自动回滚\n'
    '3. Deploy to Prod (Canary 5% → 100%)：\n'
    '   需手动审批，金丝雀监控 10 分钟后自动全量'
)

doc.add_heading('5.3 流水线 Mermaid 图', level=2)
mermaid = doc.add_paragraph()
mermaid.text = (
    '```mermaid\n'
    'graph TB\n'
    '    A[Git Push / PR] --> B{Lint}\n'
    '    B -->|pass| C[Unit Test]\n'
    '    B -->|fail| Z[中断]\n'
    '    C -->|pass| D[SAST Scan]\n'
    '    C -->|fail| Z\n'
    '    D -->|pass| E[Docker Build]\n'
    '    D -->|fail| Z\n'
    '    E -->|pass| F[Trivy Scan]\n'
    '    E -->|fail| Z\n'
    '    F -->|pass| G[Integration Test]\n'
    '    F -->|fail| Z\n'
    '    G -->|pass| H[Push Image to GHCR]\n'
    '    G -->|fail| Z\n'
    '    H --> I[Deploy to Dev]\n'
    '    I --> J{main branch?}\n'
    '    J -->|yes| K[Deploy Staging Canary 5%]\n'
    '    J -->|no| L[Done]\n'
    '    K --> M{Error < 1% & P95 < 500ms?}\n'
    '    M -->|yes| N[Canary 25% → 50% → 100%]\n'
    '    M -->|no| O[Auto Rollback]\n'
    '    N --> P{manual approve?}\n'
    '    P -->|yes| Q[Deploy Prod Canary 5%]\n'
    '    P -->|no| L\n'
    '    Q --> R{Monitor 10 min}\n'
    '    R -->|healthy| S[Promote to 100%]\n'
    '    R -->|unhealthy| O\n'
    '    S --> L\n'
    '```'
)

# ==== 第6章：渐进式发布策略 ====
doc.add_heading('6 渐进式发布策略', level=1)

doc.add_heading('6.1 金丝雀发布 (Canary Release)', level=2)
doc.add_paragraph(
    '选择金丝雀发布而非蓝绿部署，理由：\n'
    '• 金丝雀：逐步切换流量，风险可控，无需两套完整环境\n'
    '• 蓝绿：需要两套环境，资源成本高，更适合大版本升级\n\n'
    '流量切分策略：\n'
    '1. 部署 Canary 实例（1 个 Pod，占总数 5%）\n'
    '2. 通过 K8s Service Label Selector 控制流量路由\n'
    '3. 监控 5-10 分钟，验证健康指标\n'
    '4. 渐进推进：5% → 25% → 50% → 100%\n'
    '5. 每步推进前验证错误率和延迟\n'
    '6. 达到 100% 后清理 Canary 资源'
)

doc.add_heading('6.2 自动回滚策略', level=2)
doc.add_paragraph(
    '触发条件（任一满足即回滚）：\n\n'
    '• P95 延迟 > 500ms 持续 3 分钟\n'
    '• 5xx 错误率 > 1% 持续 2 分钟\n'
    '• 服务 Pod 全部 Unavailable\n\n'
    '回滚流程：\n'
    '1. 检测到异常 → Prometheus AlertManager 触发\n'
    '2. GitHub Actions 自动执行 rollback.sh\n'
    '3. Helm rollback 到上一稳定版本\n'
    '4. 验证回滚后健康检查\n'
    '5. 发送通知（GitHub Issue + 团队群消息）\n'
    '6. 记录 Postmortem'
)

doc.add_heading('6.3 一键回滚脚本', level=2)
doc.add_paragraph(
    '提供 bash scripts/rollback.sh <env> 一键回滚脚本：\n'
    '• 自动查找上一个稳定 Helm revision\n'
    '• 执行 helm rollback\n'
    '• 验证 deployment rollout status\n'
    '• 健康检查确认\n'
    '• 记录回滚事件日志'
)

# ==== 第7章：可观测性方案 ====
doc.add_heading('7 可观测性方案', level=1)

doc.add_heading('7.1 三信号架构', level=2)
doc.add_paragraph(
    '1. Logging（日志）\n'
    '   • 结构化 JSON 日志（structlog）\n'
    '   • 每条日志含 traceId 字段\n'
    '   • 日志采集：Loki + Promtail\n'
    '   • 脱敏处理：手机号、身份证自动掩码\n\n'
    '2. Metrics（指标）\n'
    '   • RED 指标：Rate / Errors / Duration\n'
    '   • USE 指标：CPU / Memory / Disk\n'
    '   • 采集：Prometheus + ServiceMonitor\n'
    '   • Dashboard：Grafana（4 个面板）\n\n'
    '3. Tracing（链路追踪）\n'
    '   • OpenTelemetry SDK 自动注入\n'
    '   • 每个请求生成唯一 traceId\n'
    '   • 后端：Tempo（OTLP gRPC）\n'
    '   • 采样策略：Dev 100% / Prod 10%'
)

doc.add_heading('7.2 Grafana Dashboard', level=2)
doc.add_paragraph(
    'Dashboard 包含 4 个面板：\n'
    '1. QPS（Queries Per Second）：按服务和路径分组的请求速率\n'
    '2. Response Latency（P50/P95/P99）：请求延迟分位数\n'
    '3. Error Rate（5xx）：错误率仪表盘\n'
    '4. Resource Usage：CPU 和 Memory 使用趋势'
)

doc.add_heading('7.3 告警规则', level=2)
doc.add_paragraph(
    '共配置 5 条告警规则：\n'
    '1. ServiceDown (Critical)：服务宕机 > 1 分钟\n'
    '2. HighErrorRate (Critical)：5xx 错误率 > 1%\n'
    '3. HighLatency (Warning)：P95 延迟 > 500ms\n'
    '4. HighMemoryUsage (Warning)：内存使用 > 85%\n'
    '5. HPANearMaxReplicas (Warning)：HPA 接近上限'
)

# ==== 第8章：安全基线 ====
doc.add_heading('8 安全基线', level=1)

doc.add_paragraph(
    '1. Secret 管理\n'
    '   • 所有敏感信息通过 GitHub Secrets + K8s Secret 注入\n'
    '   • 严禁在代码、配置文件、Dockerfile 中硬编码凭证\n'
    '   • .gitignore 排除 .env 文件\n\n'
    '2. 容器安全\n'
    '   • 多阶段构建最小化攻击面\n'
    '   • 非 root 用户运行（UID 1000）\n'
    '   • 只读根文件系统（readOnlyRootFilesystem）\n'
    '   • 丢弃所有 Linux Capabilities\n'
    '   • Trivy 扫描无 HIGH/CRITICAL 漏洞\n\n'
    '3. 代码安全\n'
    '   • Bandit SAST 扫描 Python 代码\n'
    '   • Hadolint 检查 Dockerfile 最佳实践\n'
    '   • Kube-linter 检查 K8s 清单安全性\n\n'
    '4. 供应链安全\n'
    '   • 锁定依赖版本（requirements.txt 精确版本）\n'
    '   • 生成 SBOM 软件物料清单\n'
    '   • 镜像签名（可选）'
)

# ==== 第9章：风险与缓解 ====
doc.add_heading('9 风险与缓解', level=1)

risks = [
    ('团队对 K8s/Helm 不熟悉',
     '提前学习官方文档，使用 Minikube 搭建本地实验环境；编写详细 runbook'),
    ('CI 流水线不稳定导致开发效率下降',
     'CI 超时 10 分钟自动取消；关键阶段配置缓存；失败时自动通知提交者'),
    ('金丝雀发布可能引入生产故障',
     '从 5% 极小比例开始；设置严格健康阈值；自动回滚机制兜底'),
    ('Docker 镜像 > 200 MB',
     '使用 slim 基础镜像 + 多阶段构建 + 清理缓存层'),
    ('可观测性数据量过大',
     '配置日志保留策略（7 天）；采样率根据环境调整；指标降采样'),
]

table = doc.add_table(rows=len(risks) + 1, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = '风险'
table.rows[0].cells[1].text = '缓解措施'
for i, (risk, mitigation) in enumerate(risks):
    table.rows[i + 1].cells[0].text = risk
    table.rows[i + 1].cells[1].text = mitigation

doc.add_page_break()
doc.add_paragraph()
disclaimer = doc.add_paragraph()
disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = disclaimer.add_run(
    '本人承诺所提交的实验材料系本人（团队）独立完成，对引用的他人成果均已明确标注。'
    'AI 生成内容已在附录中说明使用范围与提示词，并对其正确性负责。'
)
run.font.size = Pt(10)
run.italic = True

# 保存
output_path = os.path.join(BASE_DIR, 'D3-1_DevOps设计方案.docx')
doc.save(output_path)
print(f'✅ D3-1 generated: {output_path}')
